"""
Markdown to DOCX Converter
Converts markdown-formatted resume text into a professional DOCX document.
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def markdown_to_docx(markdown_text: str) -> io.BytesIO:
    """
    Convert markdown text to a DOCX file in memory.
    """
    document = Document()
    
    # Set margins (0.5 inch)
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Default style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Update all styles to use Times New Roman
    for s in document.styles:
        if hasattr(s, 'font'):
            s.font.name = 'Times New Roman'
    
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Headers
        if line.startswith('# '):
            # Name (H1)
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
            
        elif line.startswith('## '):
            # Section Header (H2)
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(line[3:].strip().upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Add bottom border to section header
            p_element = p._p
            pPr = p_element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '000000')
            pBdr.append(bottom)
            pPr.append(pBdr)
            
        elif line.startswith('### '):
            # Sub-header (H3) - e.g. Job Title
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.font.size = Pt(11)
            
        elif line.startswith('#### '):
            # H4
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(line[5:].strip())
            run.bold = True
            run.italic = True
            
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet point
            p = document.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(0)
            # Remove bold/markdown syntax from within the line
            content = line[2:].strip()
            
            # Handle bolding **text**
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
                    
        else:
            # Normal text
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            # Handle bolding **text**
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    # Save to memory
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream
